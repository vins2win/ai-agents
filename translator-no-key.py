import os
import docx
from transformers import MarianMTModel, MarianTokenizer
import argparse

class DocTranslator:
    def __init__(self):
        # Load languages and their codes for MarianMT
        self.language_map = {
            'german': 'de',
            'french': 'fr',
            'spanish': 'es',
            'italian': 'it',
            'dutch': 'nl',
            'polish': 'pl',
            'portuguese': 'pt',
            'russian': 'ru',
            # Add more languages as needed
        }
        
        # Initialize default model and tokenizer for English to German
        self.current_target_lang = 'de'
        self.model_name = f'Helsinki-NLP/opus-mt-en-{self.current_target_lang}'
        self.model = MarianMTModel.from_pretrained(self.model_name)
        self.tokenizer = MarianTokenizer.from_pretrained(self.model_name)
        
    def load_document(self, file_path):
        """Load a Word document and extract its text."""
        try:
            doc = docx.Document(file_path)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            return "\n".join(text)
        except Exception as e:
            return f"Error loading document: {str(e)}"
    
    def change_target_language(self, language):
        """Change the target language for translation."""
        language = language.lower()
        if language in self.language_map:
            code = self.language_map[language]
            try:
                self.current_target_lang = code
                self.model_name = f'Helsinki-NLP/opus-mt-en-{code}'
                self.model = MarianMTModel.from_pretrained(self.model_name)
                self.tokenizer = MarianTokenizer.from_pretrained(self.model_name)
                return f"Successfully switched target language to {language} ({code})"
            except Exception as e:
                return f"Error loading model for {language}: {str(e)}"
        else:
            return f"Language '{language}' not supported. Supported languages: {', '.join(self.language_map.keys())}"
    
    def translate_text(self, text):
        """Translate text to the current target language."""
        try:
            # Split text into manageable chunks (2000 characters each)
            chunks = [text[i:i+2000] for i in range(0, len(text), 2000)]
            translated_chunks = []
            
            for chunk in chunks:
                inputs = self.tokenizer(chunk, return_tensors="pt", padding=True, truncation=True, max_length=512)
                translated = self.model.generate(**inputs)
                translated_text = self.tokenizer.batch_decode(translated, skip_special_tokens=True)[0]
                translated_chunks.append(translated_text)
            
            return "".join(translated_chunks)
        except Exception as e:
            return f"Error during translation: {str(e)}"
    
    def save_translated_document(self, original_path, translated_text):
        """Save translated text to a new Word document."""
        try:
            # Create a new filename with language code
            base_name = os.path.splitext(original_path)[0]
            new_path = f"{base_name}_{self.current_target_lang}.docx"
            
            # Create a new document and add the translated text
            doc = docx.Document()
            for paragraph in translated_text.split('\n'):
                doc.add_paragraph(paragraph)
            
            doc.save(new_path)
            return f"Translated document saved as {new_path}"
        except Exception as e:
            return f"Error saving document: {str(e)}"

def main():
    parser = argparse.ArgumentParser(description='Translate Word documents to different languages')
    parser.add_argument('file_path', help='Path to the Word document to translate')
    parser.add_argument('--language', '-l', default='german', 
                        help='Target language for translation (e.g., german, french, spanish)')
    args = parser.parse_args()
    
    translator = DocTranslator()
    
    # Change language if specified
    if args.language.lower() != 'german':
        print(translator.change_target_language(args.language))
    
    # Load document
    print(f"Loading document: {args.file_path}")
    document_text = translator.load_document(args.file_path)
    
    if document_text.startswith("Error"):
        print(document_text)
        return
    
    # Translate document
    print(f"Translating document to {args.language}...")
    translated_text = translator.translate_text(document_text)
    
    if translated_text.startswith("Error"):
        print(translated_text)
        return
    
    # Save translated document
    result = translator.save_translated_document(args.file_path, translated_text)
    print(result)

if __name__ == "__main__":
    main()
