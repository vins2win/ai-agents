import os
import docx
from transformers import MarianMTModel, MarianTokenizer
from langchain.agents import Tool, AgentType, initialize_agent
from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate

class DocTranslator:
    def __init__(self):
        # Load languages and their codes for MarianMT
        self.language_map = {
            'german': 'de',
            'french': 'fr',
            'spanish': 'es',
            'italian': 'it',
            'dutch': 'nl',
            'polish': 'pl',
            'portuguese': 'pt',
            'russian': 'ru',
            # Add more languages as needed
        }
        
        # Initialize default model and tokenizer for English to German
        self.current_target_lang = 'de'
        self.model_name = f'Helsinki-NLP/opus-mt-en-{self.current_target_lang}'
        self.model = MarianMTModel.from_pretrained(self.model_name)
        self.tokenizer = MarianTokenizer.from_pretrained(self.model_name)
        
    def load_document(self, file_path):
        """Load a Word document and extract its text."""
        try:
            doc = docx.Document(file_path)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            return "\n".join(text)
        except Exception as e:
            return f"Error loading document: {str(e)}"
    
    def change_target_language(self, language):
        """Change the target language for translation."""
        language = language.lower()
        if language in self.language_map:
            code = self.language_map[language]
            try:
                self.current_target_lang = code
                self.model_name = f'Helsinki-NLP/opus-mt-en-{code}'
                self.model = MarianMTModel.from_pretrained(self.model_name)
                self.tokenizer = MarianTokenizer.from_pretrained(self.model_name)
                return f"Successfully switched target language to {language} ({code})"
            except Exception as e:
                return f"Error loading model for {language}: {str(e)}"
        else:
            return f"Language '{language}' not supported. Supported languages: {', '.join(self.language_map.keys())}"
    
    def translate_text(self, text):
        """Translate text to the current target language."""
        try:
            # Split text into manageable chunks (2000 characters each)
            chunks = [text[i:i+2000] for i in range(0, len(text), 2000)]
            translated_chunks = []
            
            for chunk in chunks:
                inputs = self.tokenizer(chunk, return_tensors="pt", padding=True, truncation=True, max_length=512)
                translated = self.model.generate(**inputs)
                translated_text = self.tokenizer.batch_decode(translated, skip_special_tokens=True)[0]
                translated_chunks.append(translated_text)
            
            return "".join(translated_chunks)
        except Exception as e:
            return f"Error during translation: {str(e)}"
    
    def save_translated_document(self, original_path, translated_text):
        """Save translated text to a new Word document."""
        try:
            # Create a new filename with language code
            base_name = os.path.splitext(original_path)[0]
            new_path = f"{base_name}_{self.current_target_lang}.docx"
            
            # Create a new document and add the translated text
            doc = docx.Document()
            for paragraph in translated_text.split('\n'):
                doc.add_paragraph(paragraph)
            
            doc.save(new_path)
            return f"Translated document saved as {new_path}"
        except Exception as e:
            return f"Error saving document: {str(e)}"

# Initialize the translator
translator = DocTranslator()

# Define tools for the agent
tools = [
    Tool(
        name="LoadDocument",
        func=translator.load_document,
        description="Load a Word document from a given file path."
    ),
    Tool(
        name="ChangeTargetLanguage",
        func=translator.change_target_language,
        description="Change the target language for translation. Input should be the language name in English (e.g., 'german', 'french')."
    ),
    Tool(
        name="TranslateText",
        func=translator.translate_text,
        description="Translate the provided text to the current target language."
    ),
    Tool(
        name="SaveTranslatedDocument",
        func=translator.save_translated_document,
        description="Save the translated text to a new Word document. Requires the original document path and the translated text."
    )
]

# Initialize the LLM
llm = ChatOpenAI(temperature=0, model_name="gpt-3.5-turbo")

# Create the agent
agent = initialize_agent(
    tools,
    llm,
    agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
    verbose=True
)

# Example prompt template for the agent
prompt_template = """
You are an AI assistant that helps with document translation. 
You can:
1. Load Word documents
2. Change the target language
3. Translate text
4. Save translated documents

The user wants: {user_input}

Think step by step about how to accomplish this task.
"""

prompt = PromptTemplate(
    input_variables=["user_input"],
    template=prompt_template,
)

translation_chain = LLMChain(llm=llm, prompt=prompt)

def process_user_request(user_input):
    """Process a user request and run the agent."""
    response = translation_chain.run(user_input=user_input)
    return agent.run(response)

# Example usage
if __name__ == "__main__":
    # Example: Translate a document to German
    result = process_user_request("Please translate the document 'sample.docx' to German and save it as a new file.")
    print(result)
    
    # Example: Change language and translate another document
    result = process_user_request("Change the language to Spanish and translate 'business_report.docx'.")
    print(result)
